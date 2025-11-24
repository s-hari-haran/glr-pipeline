#!/usr/bin/env python3
"""
Headless generator for GLR Pipeline — reproduces the app's generation flow
Creates a sample template, runs LLM extraction on the example photo report text,
asks the LLM to fill the template, and writes `Completed_GLR_Report.docx` to the workspace.
"""
import os
import sys
import json
from pathlib import Path
from dotenv import load_dotenv

# Ensure workspace root is in path for local imports
ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / 'glr_pipeline_app'))

from llm_handler import GeminiLLMHandler
from template_handler import DocxTemplateHandler
from data_mapper import DataMapper

# create a simple template if none exists
from docx import Document

def create_sample_template(path: str):
    doc = Document()
    doc.add_heading('GLR Sample Template', level=1)
    doc.add_paragraph('Insured: [INSURED_NAME]')
    doc.add_paragraph('Policy #: [POLICY_NUMBER]')
    doc.add_paragraph('Claim #: [CLAIM_NUMBER]')
    doc.add_paragraph('Date Inspected: [DATE_INSPECTED]')
    doc.add_paragraph('Address: [INSURED_H_STREET], [INSURED_H_CITY], [INSURED_H_STATE] [INSURED_H_ZIP]')
    doc.add_paragraph('Type of Loss: [TOL_CODE]')
    doc.add_paragraph('\nDamage Summary:\n[damage_summary]')
    doc.save(path)


def load_example_text():
    # Example file from the repo
    candidate = Path(ROOT) / 'Task 3 - GLR Pipeline' / 'Example 1 - USAA' / 'Input' / 'photo report.txt'
    if not candidate.exists():
        print('Example photo report not found:', candidate)
        return None
    return candidate.read_text(encoding='utf-8')


def main():
    load_dotenv()
    api_key = os.environ.get('GOOGLE_API_KEY')
    if not api_key:
        print('GOOGLE_API_KEY not set in environment or .env; aborting.')
        return 2

    example_text = load_example_text()
    if not example_text:
        print('No example photo report text available; aborting.')
        return 3

    # Create a template file in workspace
    template_path = str(Path.cwd() / 'sample_template.docx')
    create_sample_template(template_path)
    print('Created template at', template_path)

    # Initialize handlers
    llm = GeminiLLMHandler(api_key)
    print('Extracting data from photo report...')
    extracted = llm.extract_insurance_data(example_text)
    print('Extracted keys:', list(extracted.keys()))

    print('Generating narratives...')
    narratives = llm.generate_narrative(extracted)
    extracted.update(narratives)

    # Load template handler
    th = DocxTemplateHandler(template_path)

    mapper = DataMapper(extracted, th.get_placeholders())
    replacements = mapper.map_data()
    report = mapper.get_mapping_report()
    mapping_path = Path.cwd() / 'mapping_report.json'
    mapping_path.write_text(json.dumps(report, indent=2), encoding='utf-8')
    print('Mapping report written to', mapping_path)

    # Preferred flow: ask LLM to produce placeholder->value mapping, then fill original docx
    try:
        placeholders = sorted(list(th.get_placeholders()))
        print('Requesting LLM to generate placeholder mapping...')
        llm_mapping = llm.generate_placeholder_mapping(placeholders, extracted)
        print('LLM mapping sample:', {k: llm_mapping.get(k) for k in placeholders[:5]})

        # Merge with local replacements as fallback
        final_replacements = {}
        for p in placeholders:
            v = llm_mapping.get(p, "")
            if v is None or v == "":
                v = replacements.get(p, "")
            final_replacements[p] = v

        out_path = Path.cwd() / 'Completed_GLR_REPORT.docx'
        th.fill_and_save(final_replacements, str(out_path))
        print('Saved filled doc to', out_path)
    except Exception as e:
        print('LLM mapping flow failed:', e)
        print('Falling back to local replacements with python-docx fill')
        try:
            tmp_out = Path.cwd() / 'Completed_GLR_REPORT.docx'
            th.fill_and_save(replacements, str(tmp_out))
            print('Saved locally-filled doc to', tmp_out)
        except Exception as e2:
            print('Local fill also failed:', e2)
            return 4

    # Basic inspection: report paragraph count
    try:
        from docx import Document as Docx
        doc = Docx(str(Path.cwd() / 'Completed_GLR_REPORT.docx'))
        print('Generated document paragraphs count:', len(doc.paragraphs))
    except Exception as e:
        print('Could not inspect generated docx:', e)

    return 0

if __name__ == '__main__':
    sys.exit(main())
