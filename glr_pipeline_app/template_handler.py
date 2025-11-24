"""
Template Parsing and Manipulation Module
Handles identification and manipulation of .docx template placeholders
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import logging
from typing import Dict, List, Tuple, Set
from copy import deepcopy

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocxTemplateHandler:
    """Handles parsing and manipulation of .docx templates"""
    
    PLACEHOLDER_PATTERN = re.compile(r'\[([A-Z_0-9]+)\]')
    
    def __init__(self, docx_path: str):
        """
        Initialize template handler with a .docx file.
        
        Args:
            docx_path: Path to the .docx template file
        """
        self.docx_path = docx_path
        self.document = Document(docx_path)
        self.placeholders = self._extract_all_placeholders()
        logger.info(f"Template loaded with {len(self.placeholders)} unique placeholders")
    
    def _extract_all_placeholders(self) -> Set[str]:
        """
        Extract all placeholder names from the document.
        
        Returns:
            Set of placeholder names found in the document
        """
        placeholders = set()
        
        # Check paragraphs
        for paragraph in self.document.paragraphs:
            matches = self.PLACEHOLDER_PATTERN.findall(paragraph.text)
            placeholders.update(matches)
        
        # Check tables
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = self.PLACEHOLDER_PATTERN.findall(paragraph.text)
                        placeholders.update(matches)
        
        logger.info(f"Found placeholders: {sorted(placeholders)}")
        return placeholders
    
    def get_placeholders(self) -> Set[str]:
        """Get all placeholders in the template"""
        return self.placeholders
    
    def replace_text_in_paragraph(self, paragraph, replacements: Dict[str, str]) -> bool:
        """
        Replace placeholders in a paragraph while preserving formatting.
        
        Args:
            paragraph: Paragraph object from python-docx
            replacements: Dictionary of {placeholder: value}
            
        Returns:
            True if any replacements were made
        """
        changed = False
        
        for placeholder, value in replacements.items():
            placeholder_text = f"[{placeholder}]"
            
            if placeholder_text in paragraph.text:
                # Handle replacement in a way that preserves formatting
                full_text = paragraph.text
                
                if placeholder_text in full_text:
                    # Split and rebuild the paragraph
                    new_text = full_text.replace(placeholder_text, str(value or ""))
                    
                    # Clear existing runs
                    for run in paragraph.runs:
                        run.text = ""
                    
                    # Add the new text as a single run, preserving paragraph formatting
                    new_run = paragraph.add_run(new_text)
                    
                    # Try to preserve original formatting from first run if it exists
                    if paragraph.runs and len(paragraph.runs) > 1:
                        first_run = paragraph.runs[0]
                        if first_run.font.size:
                            new_run.font.size = first_run.font.size
                        if first_run.font.bold:
                            new_run.font.bold = first_run.font.bold
                    
                    changed = True
                    logger.info(f"Replaced [{placeholder}] in paragraph")
        
        return changed
    
    def fill_template(self, replacements: Dict[str, str]) -> Document:
        """
        Fill template with provided values.
        
        Args:
            replacements: Dictionary of {placeholder: value}
            
        Returns:
            Filled Document object
        """
        filled_doc = deepcopy(self.document)
        replacements_count = 0
        
        # Replace in paragraphs
        for paragraph in filled_doc.paragraphs:
            if self.replace_text_in_paragraph(paragraph, replacements):
                replacements_count += 1
        
        # Replace in tables
        for table in filled_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if self.replace_text_in_paragraph(paragraph, replacements):
                            replacements_count += 1
        
        logger.info(f"Completed {replacements_count} replacements in template")
        return filled_doc
    
    def fill_and_save(self, replacements: Dict[str, str], output_path: str) -> bool:
        """
        Fill template and save to output file.
        
        Args:
            replacements: Dictionary of {placeholder: value}
            output_path: Path where filled document should be saved
            
        Returns:
            True if successful
        """
        try:
            filled_doc = self.fill_template(replacements)
            filled_doc.save(output_path)
            logger.info(f"Filled template saved to {output_path}")
            return True
        except Exception as e:
            logger.error(f"Error saving filled template: {e}")
            raise
    
    def get_placeholder_mapping_template(self) -> Dict[str, str]:
        """
        Get a template dictionary with all placeholders for user reference.
        
        Returns:
            Dictionary with all placeholders as keys and empty strings as values
        """
        return {placeholder: "" for placeholder in sorted(self.placeholders)}

    def get_template_text(self) -> str:
        """
        Return the template content as plain text with placeholders preserved.
        This is useful when we want to generate a filled document from plain text.
        """
        parts = []
        for paragraph in self.document.paragraphs:
            parts.append(paragraph.text)

        # Include tables as text blocks
        for table in self.document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    # join cell paragraphs
                    cell_text = "\n".join([p.text for p in cell.paragraphs if p.text])
                    row_text.append(cell_text)
                parts.append(" | ".join(row_text))

        return "\n\n".join(parts)
